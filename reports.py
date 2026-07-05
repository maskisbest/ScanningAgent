"""
Report exporters for HTML, plain text and Word formats.

The Word export uses python-docx and creates a standard .docx document. This
module is original coursework code.
"""

from __future__ import annotations

import html
import re
from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.shared import Inches


SEVERITY_ORDER = ["高危", "中危", "低危", "信息"]

# banner/evidence text may contain raw bytes from non-HTTP services (e.g. MySQL
# handshake packets); strip anything that isn't valid in XML 1.0 before it goes
# into a docx run, or python-docx raises ValueError.
_XML_INVALID_RE = re.compile("[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _safe_text(value: Any) -> str:
    return _XML_INVALID_RE.sub("", str(value))


def build_text_report(summary: Dict[str, Any], author_info: Dict[str, str]) -> str:
    scan = summary["scan"]
    lines = [
        "网络漏洞扫描课程设计报告",
        "=" * 32,
        f"开发人员：{author_info.get('student_name', '')}",
        f"学号：{author_info.get('student_id', '')}",
        f"班级：{author_info.get('class_name', '')}",
        f"指导老师：{author_info.get('teacher', '')}",
        "",
        f"扫描目标：{scan['target_input']}",
        f"端口范围：{scan['ports']}",
        f"并发线程：{scan['threads']}",
        f"扫描状态：{scan['status']}",
        f"开始时间：{scan['started_at']}",
        f"结束时间：{scan.get('finished_at') or ''}",
        f"扫描时长：{scan.get('duration_seconds') or 0} 秒",
        f"目标数量：{scan.get('target_count') or 0}",
        f"开放端口：{summary['service_count']}",
        f"漏洞/风险数量：{summary['finding_count']}",
        "",
        "风险等级统计：",
    ]
    for severity in SEVERITY_ORDER:
        lines.append(f"- {severity}: {summary['by_severity'].get(severity, 0)}")
    lines.extend(["", "OWASP Top 10:2021 覆盖："])
    for name, count in summary.get("by_owasp", {}).items():
        lines.append(f"- {name}: {count}")
    lines.extend(["", "等保 2.0 技术类别覆盖："])
    for name, count in summary.get("by_dengbao", {}).items():
        lines.append(f"- {name}: {count}")
    lines.append(f"\n关联 CVE 参考数量：{summary.get('cve_count', 0)}")
    lines.extend(["", "开放服务："])
    for service in summary["services"]:
        lines.append(f"- {service['host']}:{service['port']} {service['service']} {service.get('banner') or ''}".strip())
    lines.extend(["", "漏洞明细："])
    for idx, finding in enumerate(summary["findings"], 1):
        cve_text = "、".join(f"{c['id']}({c['summary']})" for c in finding.get("cve_refs") or []) or "无"
        lines.extend(
            [
                f"{idx}. [{finding['severity']}] {finding['title']}",
                f"   目标：{finding['host']}:{finding.get('port') or '-'}",
                f"   类型：{finding['category']} / {finding['rule_id']}",
                f"   合规映射：{finding.get('owasp_category') or '-'} / {finding.get('dengbao_category') or '-'}",
                f"   关联 CVE：{cve_text}",
                f"   证据：{finding.get('evidence') or ''}",
                f"   建议：{finding.get('recommendation') or ''}",
                f"   来源：{finding.get('source') or ''}",
            ]
        )
    return "\n".join(lines)


def build_html_report(summary: Dict[str, Any], author_info: Dict[str, str]) -> str:
    scan = summary["scan"]
    severity_cards = "".join(
        f"<div class='card'><strong>{html.escape(sev)}</strong><span>{summary['by_severity'].get(sev, 0)}</span></div>"
        for sev in SEVERITY_ORDER
    )
    service_rows = "".join(
        "<tr>"
        f"<td>{html.escape(str(item['host']))}</td>"
        f"<td>{item['port']}</td>"
        f"<td>{html.escape(item['service'])}</td>"
        f"<td>{html.escape(item.get('banner') or '')}</td>"
        "</tr>"
        for item in summary["services"]
    )
    def _compliance_html(item: Dict[str, Any]) -> str:
        parts = []
        if item.get("owasp_category"):
            parts.append(html.escape(item["owasp_category"]))
        if item.get("dengbao_category"):
            parts.append(html.escape(item["dengbao_category"]))
        for cve in item.get("cve_refs") or []:
            parts.append(f"<span title='{html.escape(cve['summary'])}'>{html.escape(cve['id'])}</span>")
        return "<br>".join(parts) or "-"

    finding_rows = "".join(
        "<tr>"
        f"<td>{html.escape(item['severity'])}</td>"
        f"<td>{html.escape(item['host'])}:{item.get('port') or '-'}</td>"
        f"<td>{html.escape(item['title'])}</td>"
        f"<td>{html.escape(item['category'])}</td>"
        f"<td>{_compliance_html(item)}</td>"
        f"<td>{html.escape(item.get('evidence') or '')}</td>"
        f"<td>{html.escape(item.get('recommendation') or '')}</td>"
        "</tr>"
        for item in summary["findings"]
    )
    owasp_rows = "".join(
        f"<tr><td>{html.escape(name)}</td><td>{count}</td></tr>"
        for name, count in summary.get("by_owasp", {}).items()
    )
    dengbao_rows = "".join(
        f"<tr><td>{html.escape(name)}</td><td>{count}</td></tr>"
        for name, count in summary.get("by_dengbao", {}).items()
    )
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>网络漏洞扫描报告 #{scan['id']}</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; margin: 32px; color: #1f2937; }}
    h1, h2 {{ color: #111827; }}
    table {{ width: 100%; border-collapse: collapse; margin: 12px 0 24px; }}
    th, td {{ border: 1px solid #d1d5db; padding: 8px; text-align: left; vertical-align: top; }}
    th {{ background: #f3f4f6; }}
    .grid {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; }}
    .card {{ border: 1px solid #d1d5db; border-radius: 8px; padding: 14px; }}
    .card span {{ display: block; font-size: 28px; margin-top: 8px; }}
  </style>
</head>
<body>
  <h1>网络漏洞扫描课程设计报告</h1>
  <p>开发人员：{html.escape(author_info.get('student_name', ''))}　
     学号：{html.escape(author_info.get('student_id', ''))}　
     班级：{html.escape(author_info.get('class_name', ''))}　
     指导老师：{html.escape(author_info.get('teacher', ''))}</p>
  <h2>扫描概况</h2>
  <p>目标：{html.escape(scan['target_input'])}</p>
  <p>端口：{html.escape(scan['ports'])}；线程：{scan['threads']}；时长：{scan.get('duration_seconds') or 0} 秒；目标数量：{scan.get('target_count') or 0}</p>
  <div class="grid">{severity_cards}</div>
  <h2>合规覆盖统计</h2>
  <p>关联 CVE 参考数量：{summary.get('cve_count', 0)}</p>
  <table><thead><tr><th>OWASP Top 10:2021</th><th>命中数</th></tr></thead><tbody>{owasp_rows}</tbody></table>
  <table><thead><tr><th>等保 2.0 技术类别</th><th>命中数</th></tr></thead><tbody>{dengbao_rows}</tbody></table>
  <h2>开放服务</h2>
  <table><thead><tr><th>主机</th><th>端口</th><th>服务</th><th>Banner</th></tr></thead><tbody>{service_rows}</tbody></table>
  <h2>漏洞明细</h2>
  <table><thead><tr><th>等级</th><th>目标</th><th>名称</th><th>类型</th><th>合规映射 / CVE</th><th>证据</th><th>建议</th></tr></thead><tbody>{finding_rows}</tbody></table>
</body>
</html>"""


def build_docx_report(summary: Dict[str, Any], author_info: Dict[str, str]) -> BytesIO:
    scan = summary["scan"]
    doc = Document()
    doc.add_heading("网络漏洞扫描课程设计报告", 0)
    doc.add_paragraph(f"开发人员：{author_info.get('student_name', '')}")
    doc.add_paragraph(f"学号：{author_info.get('student_id', '')}")
    doc.add_paragraph(f"班级：{author_info.get('class_name', '')}")
    doc.add_paragraph(f"指导老师：{author_info.get('teacher', '')}")
    doc.add_heading("扫描概况", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    overview = [
        ("扫描目标", scan["target_input"]),
        ("端口范围", scan["ports"]),
        ("并发线程", str(scan["threads"])),
        ("扫描状态", scan["status"]),
        ("开始时间", scan["started_at"]),
        ("结束时间", scan.get("finished_at") or ""),
        ("扫描时长", f"{scan.get('duration_seconds') or 0} 秒"),
        ("目标数量", str(scan.get("target_count") or 0)),
        ("开放端口", str(summary["service_count"])),
        ("漏洞/风险数量", str(summary["finding_count"])),
    ]
    for key, value in overview:
        cells = table.add_row().cells
        cells[0].text = _safe_text(key)
        cells[1].text = _safe_text(value)
    doc.add_heading("风险等级统计", level=1)
    sev_table = doc.add_table(rows=1, cols=2)
    sev_table.style = "Table Grid"
    sev_table.rows[0].cells[0].text = "等级"
    sev_table.rows[0].cells[1].text = "数量"
    for severity in SEVERITY_ORDER:
        cells = sev_table.add_row().cells
        cells[0].text = severity
        cells[1].text = str(summary["by_severity"].get(severity, 0))
    doc.add_heading("合规覆盖统计", level=1)
    doc.add_paragraph(f"关联 CVE 参考数量：{summary.get('cve_count', 0)}")
    owasp_table = doc.add_table(rows=1, cols=2)
    owasp_table.style = "Table Grid"
    owasp_table.rows[0].cells[0].text = "OWASP Top 10:2021"
    owasp_table.rows[0].cells[1].text = "命中数"
    for name, count in summary.get("by_owasp", {}).items():
        cells = owasp_table.add_row().cells
        cells[0].text = _safe_text(name)
        cells[1].text = str(count)
    dengbao_table = doc.add_table(rows=1, cols=2)
    dengbao_table.style = "Table Grid"
    dengbao_table.rows[0].cells[0].text = "等保 2.0 技术类别"
    dengbao_table.rows[0].cells[1].text = "命中数"
    for name, count in summary.get("by_dengbao", {}).items():
        cells = dengbao_table.add_row().cells
        cells[0].text = _safe_text(name)
        cells[1].text = str(count)
    doc.add_heading("开放服务", level=1)
    service_table = doc.add_table(rows=1, cols=4)
    service_table.style = "Table Grid"
    for idx, title in enumerate(["主机", "端口", "服务", "Banner"]):
        service_table.rows[0].cells[idx].text = title
    for service in summary["services"]:
        cells = service_table.add_row().cells
        cells[0].text = _safe_text(service["host"])
        cells[1].text = str(service["port"])
        cells[2].text = _safe_text(service["service"])
        cells[3].text = _safe_text(service.get("banner") or "")
    doc.add_heading("漏洞明细", level=1)
    for idx, item in enumerate(summary["findings"], 1):
        doc.add_heading(_safe_text(f"{idx}. [{item['severity']}] {item['title']}"), level=2)
        doc.add_paragraph(_safe_text(f"目标：{item['host']}:{item.get('port') or '-'}"))
        doc.add_paragraph(_safe_text(f"类型：{item['category']} / {item['rule_id']}"))
        cve_text = "、".join(f"{c['id']}({c['summary']})" for c in item.get("cve_refs") or []) or "无"
        doc.add_paragraph(_safe_text(f"合规映射：{item.get('owasp_category') or '-'} / {item.get('dengbao_category') or '-'}"))
        doc.add_paragraph(_safe_text(f"关联 CVE：{cve_text}"))
        doc.add_paragraph(_safe_text(f"证据：{item.get('evidence') or ''}"))
        doc.add_paragraph(_safe_text(f"建议：{item.get('recommendation') or ''}"))
        doc.add_paragraph(_safe_text(f"来源：{item.get('source') or ''}"))
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
